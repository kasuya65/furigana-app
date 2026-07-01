import copy
import io
from lxml import etree

import pandas as pd
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pykakasi

st.set_page_config(page_title="ふりがな自動付与", page_icon="🈺", layout="wide")
st.title("🈺 ふりがな自動付与")
st.markdown("Wordファイルをアップロードすると、漢字部分にふりがな（ルビ）を自動で付与します。")

_kks = pykakasi.kakasi()

_RUBY_MIN_LINE = 480  # twips (24pt): 8pt ルビ + 12pt 本文 + 余白


def _has_kanji(text: str) -> bool:
    return any('一' <= c <= '鿿' or '㐀' <= c <= '䶿' for c in text)


def _tokenize(text: str, overrides: dict | None = None) -> list[dict]:
    tokens = _kks.convert(text)
    result = []
    for t in tokens:
        orig = t['orig']
        hira = t['hira']
        if overrides and orig in overrides:
            hira = overrides[orig]
        result.append({'orig': orig, 'hira': hira, 'has_kanji': _has_kanji(orig)})
    return result


def _make_ruby_element(orig: str, ruby: str, rpr_elem=None) -> etree._Element:
    ruby_el = OxmlElement('w:ruby')

    rubyPr = OxmlElement('w:rubyPr')
    for tag, val in [
        ('w:rubyAlign', 'distributeSpace'),
        ('w:hps', '16'),
        ('w:hpsRaise', '25'),
        ('w:hpsBaseText', '24'),
        ('w:lid', 'ja-JP'),
    ]:
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        rubyPr.append(el)
    ruby_el.append(rubyPr)

    rt = OxmlElement('w:rt')
    rt_r = OxmlElement('w:r')
    rt_rPr = OxmlElement('w:rPr')
    for tag, val in [('w:sz', '16'), ('w:szCs', '16')]:
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        rt_rPr.append(el)
    rt_r.append(rt_rPr)
    rt_t = OxmlElement('w:t')
    rt_t.text = ruby
    if ruby and (ruby[0] == ' ' or ruby[-1] == ' '):
        rt_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    rt_r.append(rt_t)
    rt.append(rt_r)
    ruby_el.append(rt)

    rubyBase = OxmlElement('w:rubyBase')
    base_r = OxmlElement('w:r')
    if rpr_elem is not None:
        base_r.append(copy.deepcopy(rpr_elem))
    base_t = OxmlElement('w:t')
    base_t.text = orig
    if orig and (orig[0] == ' ' or orig[-1] == ' '):
        base_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    base_r.append(base_t)
    rubyBase.append(base_r)
    ruby_el.append(rubyBase)

    return ruby_el


def _make_plain_run(text: str, rpr_elem=None) -> etree._Element:
    r = OxmlElement('w:r')
    if rpr_elem is not None:
        r.append(copy.deepcopy(rpr_elem))
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r


def _process_run(run_elem: etree._Element, overrides: dict) -> list[etree._Element]:
    t_elem = run_elem.find(qn('w:t'))
    if t_elem is None or not t_elem.text:
        return [run_elem]
    text = t_elem.text
    if not _has_kanji(text):
        return [run_elem]
    rpr_elem = run_elem.find(qn('w:rPr'))
    return [
        _make_ruby_element(tok['orig'], tok['hira'], rpr_elem) if tok['has_kanji']
        else _make_plain_run(tok['orig'], rpr_elem)
        for tok in _tokenize(text, overrides)
    ]


def _adjust_spacing_for_ruby(para_elem: etree._Element) -> None:
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    line_rule = spacing.get(qn('w:lineRule'), 'auto')
    line_val_str = spacing.get(qn('w:line'))
    if line_rule == 'atLeast' and line_val_str and int(line_val_str) >= _RUBY_MIN_LINE:
        return
    spacing.set(qn('w:lineRule'), 'atLeast')
    spacing.set(qn('w:line'), str(_RUBY_MIN_LINE))


def _process_paragraph(para_elem: etree._Element, overrides: dict) -> bool:
    runs = para_elem.findall(qn('w:r'))
    ruby_added = False
    for run in runs:
        replacements = _process_run(run, overrides)
        if len(replacements) == 1 and replacements[0] is run:
            continue
        ruby_added = True
        parent = run.getparent()
        idx = list(parent).index(run)
        parent.remove(run)
        for i, el in enumerate(replacements):
            parent.insert(idx + i, el)
    return ruby_added


def extract_readings(docx_bytes: bytes) -> dict[str, str]:
    doc = Document(io.BytesIO(docx_bytes))
    seen: dict[str, str] = {}

    def process_text(text: str) -> None:
        if not _has_kanji(text):
            return
        for tok in _tokenize(text):
            if tok['has_kanji'] and tok['orig'] not in seen:
                seen[tok['orig']] = tok['hira']

    for para in doc.paragraphs:
        for run in para.runs:
            process_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        process_text(run.text)
    return seen


def add_furigana(docx_bytes: bytes, overrides: dict) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    for para in doc.paragraphs:
        if _process_paragraph(para._p, overrides):
            _adjust_spacing_for_ruby(para._p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _process_paragraph(para._p, overrides):
                        _adjust_spacing_for_ruby(para._p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_preview_html(docx_bytes: bytes, overrides: dict) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    for para in doc.paragraphs:
        if not para.text.strip():
            lines.append('<p style="margin:0.2em 0">&nbsp;</p>')
            continue
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if _has_kanji(text):
                for tok in _tokenize(text, overrides):
                    if tok['has_kanji']:
                        parts.append(f'<ruby>{tok["orig"]}<rt>{tok["hira"]}</rt></ruby>')
                    else:
                        parts.append(tok['orig'])
            else:
                parts.append(text)
        lines.append(f'<p style="margin:0.6em 0;line-height:2.2">{"".join(parts)}</p>')
    return '<div style="font-size:16px;font-family:serif;padding:1em;border:1px solid #ddd;border-radius:6px">' + ''.join(lines) + '</div>'


# ── UI ──────────────────────────────────────────────────────────────

uploaded = st.file_uploader("Wordファイルをアップロード (.docx)", type=["docx"])

if uploaded:
    file_bytes = uploaded.read()

    if 'furigana_source' not in st.session_state or st.session_state.furigana_source != uploaded.name:
        st.session_state.furigana_source = uploaded.name
        st.session_state.furigana_readings = None

    st.success(f"アップロード完了：{uploaded.name}")

    if st.button("ふりがなを解析する", type="primary"):
        with st.spinner("解析中..."):
            readings = extract_readings(file_bytes)
        st.session_state.furigana_readings = readings
        st.session_state.furigana_file_bytes = file_bytes

if st.session_state.get('furigana_readings'):
    readings: dict = st.session_state.furigana_readings
    src_bytes: bytes = st.session_state.furigana_file_bytes

    st.markdown("### ふりがな一覧・修正")
    st.caption("辞書ベース（pykakasi）で自動検出した読みです。間違いがあればふりがな欄を直接編集してください。")

    df = pd.DataFrame(
        [{'漢字': k, 'ふりがな': v} for k, v in readings.items()]
    )
    edited = st.data_editor(
        df,
        use_container_width=True,
        num_rows="fixed",
        column_config={
            '漢字': st.column_config.TextColumn('漢字', disabled=True),
            'ふりがな': st.column_config.TextColumn('ふりがな（編集可）'),
        },
        height=min(400, 45 + len(df) * 35),
        key="furigana_editor",
    )

    overrides = dict(zip(edited['漢字'], edited['ふりがな']))

    col1, col2 = st.columns([1, 3])
    with col1:
        preview_clicked = st.button("🔍 プレビュー更新", use_container_width=True)
    with col2:
        generate_clicked = st.button("📄 Wordを生成してダウンロード", type="primary", use_container_width=True)

    if preview_clicked or 'furigana_preview_html' in st.session_state:
        if preview_clicked:
            st.session_state.furigana_preview_html = build_preview_html(src_bytes, overrides)
        st.markdown("### プレビュー")
        st.html(st.session_state.furigana_preview_html)

    if generate_clicked:
        with st.spinner("Word生成中..."):
            try:
                result_bytes = add_furigana(src_bytes, overrides)
            except Exception as e:
                import traceback
                st.error(f"エラーが発生しました: {e}")
                with st.expander("詳細"):
                    st.code(traceback.format_exc())
                st.stop()

        out_name = st.session_state.furigana_source.replace(".docx", "_ふりがな.docx")
        st.download_button(
            label="📥 ダウンロード",
            data=result_bytes,
            file_name=out_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
